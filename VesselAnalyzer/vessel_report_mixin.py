"""
vessel_report_mixin.py — Report generation (Word/PDF) and CSV export.

Part of the VesselAnalyzer mixin decomposition.
These methods are mixed into VesselAnalyzerLogic via multiple inheritance:

    class VesselAnalyzerLogic(
        ReportMixin,
        ...
        ScriptedLoadableModuleLogic,
    ): ...

All methods use ``self`` normally — no changes to call sites required.
"""
# ruff: noqa  (this file is auto-extracted; formatting is inherited from VesselAnalyzer.py)


class ReportMixin:
    """Mixin: Report generation (Word/PDF) and CSV export."""

    def generateReport(self, path):
        """Generate an AI-powered clinical DOCX matching VesselMeasurementReport style."""
        import os, sys, subprocess, json, urllib.request

        # ── ensure python-docx ────────────────────────────────────────────
        try:
            from docx import Document
        except ImportError:
            subprocess.run(
                [
                    sys.executable,
                    "-m",
                    "pip",
                    "install",
                    "python-docx",
                    "--break-system-packages",
                    "-q",
                ],
                check=True,
            )
            from docx import Document

        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # ── style constants (match reference doc) ─────────────────────────
        C_NAVY = RGBColor(0x1F, 0x4E, 0x79)
        C_BLUE = RGBColor(0x2E, 0x75, 0xB6)
        C_GREY = RGBColor(0x44, 0x44, 0x44)
        C_AMBER = RGBColor(0x7D, 0x4E, 0x00)
        C_RED = RGBColor(0x8B, 0x00, 0x00)
        C_SILVER = RGBColor(0xAA, 0xAA, 0xAA)
        C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        C_BLACK = RGBColor(0x00, 0x00, 0x00)

        HX_NAVY = "1F4E79"
        HX_LBLUE = "D6E4F0"
        HX_LGREY = "F2F2F2"
        HX_AMBER = "FFF2CC"
        HX_LGRN = "E8F5E9"
        HX_WHITE = "FFFFFF"
        HX_RED = "FFE0E0"

        # ── helpers ──────────────────────────────────────────────────────
        def shd(cell, fill):
            tc = cell._tc
            pr = tc.get_or_add_tcPr()
            s = OxmlElement("w:shd")
            s.set(qn("w:val"), "clear")
            s.set(qn("w:color"), "auto")
            s.set(qn("w:fill"), fill)
            pr.append(s)

        def borders(cell, color="CCCCCC"):
            tc = cell._tc
            pr = tc.get_or_add_tcPr()
            tb = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), color)
                tb.append(b)
            pr.append(tb)

        def margins(cell, top=80, bottom=80, left=120, right=120):
            tc = cell._tc
            pr = tc.get_or_add_tcPr()
            m = OxmlElement("w:tcMar")
            for side, val in (
                ("top", top),
                ("left", left),
                ("bottom", bottom),
                ("right", right),
            ):
                e = OxmlElement(f"w:{side}")
                e.set(qn("w:w"), str(val))
                e.set(qn("w:type"), "dxa")
                m.append(e)
            pr.append(m)

        def cell_fmt(
            cell,
            text,
            bold=False,
            color=C_BLACK,
            size=10,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            italic=False,
            bg=None,
            border_color="CCCCCC",
        ):
            if bg:
                shd(cell, bg)
            borders(cell, border_color)
            margins(cell)
            p = cell.paragraphs[0]
            p.clear()
            p.alignment = align
            run = p.add_run(str(text))
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.color.rgb = color

        def hdr_cell(cell, text, w_cm):
            cell.width = Cm(w_cm)
            cell_fmt(
                cell,
                text,
                bold=True,
                color=C_WHITE,
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bg=HX_NAVY,
            )

        def add_para(
            doc,
            text,
            size=10,
            color=C_BLACK,
            bold=False,
            italic=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=0,
            space_after=6,
        ):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.color.rgb = color
            return p

        def h1(doc, text):
            p = doc.add_heading("", level=1)
            p.clear()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = C_NAVY
            return p

        def h2(doc, text):
            p = doc.add_heading("", level=2)
            p.clear()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = C_BLUE
            return p

        def detail_table(doc, rows_data):
            tbl = doc.add_table(rows=0, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for label, value, val_bg in rows_data:
                row = tbl.add_row()
                lc, vc = row.cells[0], row.cells[1]
                lc.width = Cm(4.5)
                vc.width = Cm(12)
                cell_fmt(lc, label, bold=True, size=10, bg=HX_LGREY)
                cell_fmt(vc, value, size=10, bg=val_bg)

        # ── collect branch stats ──────────────────────────────────────────
        labels = [self.getBranchDisplayName(i) for i in range(len(self.branches) + 5)]
        bdata = []
        for i in range(self.getNumBranches()):
            s = self.getBranchStats(i)
            si, ei = self.branches[i]
            zv = [self.points[j][2] for j in range(si, ei)]
            flag = s["avg"] > 0 and (
                s["max"] > s["avg"] * 1.5
                or (s["min"] > 0 and s["min"] < s["avg"] * 0.3)
            )
            meta_i = getattr(self, "branchMeta", {}).get(i, {})
            bdata.append(
                {
                    "name": labels[i] if i < len(labels) else f"Branch {i+1}",
                    "pts": ei - si,
                    "zTop": round(max(zv), 1),
                    "zBot": round(min(zv), 1),
                    "minD": round(s["min"], 2),
                    "maxD": round(s["max"], 2),
                    "avgD": round(s["avg"], 2),
                    "length": round(s["length"], 1),
                    "flag": flag,
                    "role": meta_i.get("role", ""),
                    "angle_deg": meta_i.get("angle_deg"),
                    "lateral": meta_i.get("lateral_label", ""),  # 'Left'/'Right'/''
                    "ortho_score": meta_i.get("ortho_score"),  # renal only
                    "tip_lat_mm": meta_i.get("tip_lateral_mm"),
                    "proxD": meta_i.get("proxD"),  # post-flare proximal diameter
                    "distD": meta_i.get("distD"),  # distal diameter (sizing anchor)
                    "ost_conf": meta_i.get(
                        "ostium_confidence"
                    ),  # confidence dict or None
                }
            )

        # ── Collect stent plan context for report ─────────────────────────
        stent_plan = getattr(self, "stentPlan", None)
        stent_ctx = ""
        findings_ctx = ""
        if stent_plan:
            stype = stent_plan.get("type", "")
            pd = stent_plan.get("proxDiam", 0)
            dd = stent_plan.get("distDiam", 0)
            ln = stent_plan.get("length", 0)
            stent_ctx = (
                f"Planned stent: {stype}, proximal Ø{pd:.1f}mm, distal Ø{dd:.1f}mm, "
                f"length {ln:.0f}mm. "
            )
            if "Kissing" in stype or "Bifurcated" in stype:
                r_l = stent_plan.get("r_left", 0) * 2
                r_r = stent_plan.get("r_right", 0) * 2
                r_t = stent_plan.get("r_trunk", 0) * 2
                stent_ctx += (
                    f"Y/Kissing configuration: left limb Ø{r_l:.1f}mm, "
                    f"right limb Ø{r_r:.1f}mm, trunk Ø{r_t:.1f}mm. "
                    f"Note: residual minor axis of 8-10mm at bifurcation confluence "
                    f"is the expected mechanical floor of bifurcation anatomy — "
                    f"not a procedural failure."
                )
        findings = getattr(self, "findings", [])
        if findings:
            f_lines = []
            for f in findings[:5]:
                f_lines.append(
                    f"{f['type']} at {self.getBranchDisplayName(f['branchIdx'])}: "
                    f"Ø{f['value']:.1f}mm (ratio {f['ratio']:.1f}x)"
                )
            findings_ctx = "Detected findings: " + "; ".join(f_lines) + ". "

        # ── call Claude API ───────────────────────────────────────────────
        print("[VesselAnalyzer] Calling Claude API for AI clinical narrative...")
        meas_summary = "\n".join(
            f"- {b['name']}: length={b['length']}mm, min={b['minD']}mm, "
            f"max={b['maxD']}mm, avg={b['avgD']}mm"
            + (f", angle={b['angle_deg']:.1f}°" if b["angle_deg"] is not None else "")
            + (
                f", ortho={b['ortho_score']:.2f}"
                if b["ortho_score"] is not None
                else ""
            )
            + (
                f", ostium_conf={b['ost_conf']['grade']}(eff={b['ost_conf'].get('effective_score', b['ost_conf']['score']):.2f})"
                if b["ost_conf"]
                else ""
            )
            for b in bdata
        )
        n_br = len(bdata)
        vessel_desc = (
            f"A bifurcated vessel with {n_br} branch{'es' if n_br != 1 else ''}. "
            f"Trunk (Branch 1): avg {bdata[0]['avgD']}mm, length {bdata[0]['length']}mm. "
            if n_br > 0
            else "Vessel measurements below."
        )

        # Build iliac context string from labeled main branches
        _iliac_lines = []
        for _bd in bdata:
            if _bd["lateral"] in ("Left", "Right"):
                _al = (
                    f", angle={_bd['angle_deg']:.1f}°"
                    if _bd["angle_deg"] is not None
                    else ""
                )
                _iliac_lines.append(
                    f"{_bd['lateral']} Iliac: length={_bd['length']}mm, "
                    f"avg Ø{_bd['avgD']}mm, min Ø{_bd['minD']}mm{_al}"
                )
        iliac_ctx = (
            "Iliac branches: " + "; ".join(_iliac_lines) + ". " if _iliac_lines else ""
        )

        # Build renal context string with ortho scores
        _renal_lines = []
        for _bd in bdata:
            if _bd["role"] == "renal_vein":  # renal_fragment excluded intentionally
                _os = _bd["ortho_score"] if _bd["ortho_score"] is not None else "N/A"
                _lat = _bd["tip_lat_mm"] if _bd["tip_lat_mm"] is not None else "N/A"
                _al = (
                    f"{_bd['angle_deg']:.1f}°"
                    if _bd["angle_deg"] is not None
                    else "N/A"
                )
                # Use distal diameter as sizing anchor — it is the most reliable
                # surface-distance estimate (free of proximal flare/junction bias).
                _dist_d = _bd.get("distD")
                _prox_d = _bd.get("proxD")
                _diam_str = (
                    f"distal Ø{_dist_d:.1f}mm (sizing), prox Ø{_prox_d:.1f}mm"
                    if _dist_d is not None and _prox_d is not None
                    else f"avg Ø{_bd['avgD']}mm"
                )
                _renal_lines.append(
                    f"{_bd['name']}: length={_bd['length']}mm, "
                    f"{_diam_str}, angle={_al}, "
                    f"lateral offset={_lat}mm, ortho_score={_os:.2f}"
                    if isinstance(_os, float)
                    else f"{_bd['name']}: length={_bd['length']}mm, avg Ø{_bd['avgD']}mm"
                )
        renal_ctx = (
            "Renal/side branches: " + "; ".join(_renal_lines) + ". "
            if _renal_lines
            else ""
        )
        prompt = (
            "You are a vascular radiologist writing a structured pre-procedural report "
            "for endovascular stenting planning.\n\n"
            f"Vessel description: {vessel_desc}\n"
            f"{iliac_ctx}"
            f"{renal_ctx}"
            f"{findings_ctx}"
            f"{stent_ctx}\n"
            "Centerline measurements from CT (VMTK automated analysis):\n"
            f"{meas_summary}\n\n"
            "Note: diameter method is closest surface point (may slightly underestimate). "
            "Endpoint artifacts may cause artificially low minimum values — "
            "verify before clinical use. "
            "Ortho score (0–1): combined angle + lateral offset metric; "
            ">0.65 = strong lateral branch, <0.5 = soft/curved branch.\n\n"
            "Write EXACTLY these 5 sections with these EXACT headings:\n"
            "CLINICAL SUMMARY\n"
            "FINDINGS\n"
            "STENT SIZING RECOMMENDATIONS\n"
            "LIMITATIONS\n"
            "CONCLUSION\n\n"
            "Rules: formal radiology language, 2-4 sentences per section, "
            "reference specific measurements, no markdown, no bullet points, plain paragraphs only."
        )

        ai = {
            "summary": "",
            "findings": "",
            "stent": "",
            "limitations": "",
            "conclusion": "",
        }
        try:
            payload = json.dumps(
                {
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 1000,
                    "messages": [{"role": "user", "content": prompt}],
                }
            ).encode()
            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages",
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=30) as resp:
                raw = json.loads(resp.read())["content"][0]["text"]

            # Parse sections
            section_map = {
                "CLINICAL SUMMARY": "summary",
                "FINDINGS": "findings",
                "STENT SIZING RECOMMENDATIONS": "stent",
                "LIMITATIONS": "limitations",
                "CONCLUSION": "conclusion",
            }
            cur_key = None
            cur_lines = []
            for line in raw.split("\n"):
                stripped = line.strip()
                matched = False
                for heading, key in section_map.items():
                    if stripped.upper().startswith(heading):
                        if cur_key:
                            ai[cur_key] = " ".join(cur_lines).strip()
                        cur_key = key
                        cur_lines = []
                        matched = True
                        break
                if not matched and stripped:
                    cur_lines.append(stripped)
            if cur_key:
                ai[cur_key] = " ".join(cur_lines).strip()

        except Exception as e:
            print(f"[VesselAnalyzer] Claude API error: {e}")
            ai["summary"] = (
                "AI narrative unavailable. See quantitative measurements below. "
                "Manual clinical review required."
            )

        # ── build DOCX ────────────────────────────────────────────────────
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin = Cm(2.5)
            sec.right_margin = Cm(2.5)

        # Set default font
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)

        # ── TITLE ──
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_after = Pt(4)
        tr = tp.add_run("VESSEL MEASUREMENT REPORT")
        tr.bold = True
        tr.font.size = Pt(20)
        tr.font.color.rgb = C_NAVY

        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(12)
        sr = sp.add_run(
            "Pre-Procedural Endovascular Planning — AI-Assisted Centerline Analysis"
        )
        sr.italic = True
        sr.font.size = Pt(11)
        sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # ── PATIENT INFO TABLE ──
        pt_tbl = doc.add_table(rows=1, cols=2)
        pt_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        lc = pt_tbl.rows[0].cells[0]
        rc = pt_tbl.rows[0].cells[1]
        for c in (lc, rc):
            shd(c, HX_LBLUE)
            borders(c, "AAAAAA")
            margins(c, top=120, bottom=120, left=160, right=160)
        lc.width = Cm(8)
        rc.width = Cm(8)
        _vtype_label = getattr(self, "vesselType", "arterial")
        _vtype_display = "Venous" if _vtype_label == "venous" else "Arterial"
        _modality = "CT Venography" if _vtype_label == "venous" else "CT Angiography"
        for c, lines in (
            (
                lc,
                [
                    ("Patient:", "____________________________"),
                    ("DOB:", "____________________________"),
                ],
            ),
            (
                rc,
                [
                    (
                        "Study date:",
                        __import__("datetime").datetime.now().strftime("%d/%m/%Y"),
                    ),
                    (
                        "Report generated:",
                        __import__("datetime")
                        .datetime.now()
                        .strftime("%d/%m/%Y  %H:%M"),
                    ),
                    ("Modality:", _modality),
                    ("Vessel type:", _vtype_display),
                ],
            ),
        ):
            for label, value in lines:
                p = c.add_paragraph()
                r1 = p.add_run(label)
                r1.bold = True
                r1.font.size = Pt(10)
                r2 = p.add_run("  " + value)
                r2.font.size = Pt(10)
        # Remove first empty paragraph in each cell
        for c in (lc, rc):
            p = c.paragraphs[0]
            if not p.text.strip():
                p._element.getparent().remove(p._element)

        doc.add_paragraph()

        # ── SECTION 1: CLINICAL SUMMARY (AI) ──
        h1(doc, "1. Clinical Summary")
        if ai["summary"]:
            add_para(doc, ai["summary"], size=10)
        doc.add_paragraph()

        # ── SECTION 2: BRANCH SUMMARY TABLE ──
        h1(doc, "2. Branch Summary")
        add_para(
            doc,
            "All measurements derived from VMTK centerline analysis. "
            "Diameters computed via closest surface point method. "
            "All values in millimetres.",
            size=9,
            color=C_GREY,
            italic=True,
        )

        col_w = [5.0, 2.0, 3.2, 1.8, 1.8, 1.8, 1.8]
        hdrs = [
            "Branch / Vessel",
            "Length (mm)",
            "Z Top → Bot",
            "Min Ø (mm)",
            "Max Ø (mm)",
            "Avg Ø (mm)",
            "Status",
        ]
        sum_tbl = doc.add_table(rows=1, cols=len(hdrs))
        sum_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (h, w) in enumerate(zip(hdrs, col_w)):
            hdr_cell(sum_tbl.rows[0].cells[i], h, w)

        for i, b in enumerate(bdata):
            row = sum_tbl.add_row()
            bg = HX_LBLUE if i < 2 else HX_WHITE
            status_bg = HX_AMBER if b["flag"] else HX_LGRN
            status_txt = "⚠ Review" if b["flag"] else "✓"
            status_col = C_AMBER if b["flag"] else RGBColor(0x1B, 0x5E, 0x20)
            vals = [
                (b["name"], True, C_BLACK, bg, WD_ALIGN_PARAGRAPH.LEFT),
                (f"{b['length']:.1f}", False, C_BLACK, bg, WD_ALIGN_PARAGRAPH.CENTER),
                (
                    f"{b['zTop']:.0f} → {b['zBot']:.0f}",
                    False,
                    C_BLACK,
                    bg,
                    WD_ALIGN_PARAGRAPH.CENTER,
                ),
                (f"{b['minD']:.2f}", False, C_BLACK, bg, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{b['maxD']:.2f}", True, C_BLACK, bg, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{b['avgD']:.2f}", False, C_BLACK, bg, WD_ALIGN_PARAGRAPH.CENTER),
                (status_txt, True, status_col, status_bg, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            for j, (txt, bld, col, cbg, aln) in enumerate(vals):
                c = row.cells[j]
                c.width = Cm(col_w[j])
                cell_fmt(c, txt, bold=bld, color=col, size=10, align=aln, bg=cbg)

        doc.add_paragraph()

        # ── SECTION 3: AI FINDINGS ──
        h1(doc, "3. Findings by Vessel")
        if ai["findings"]:
            add_para(doc, ai["findings"], size=10)
        else:
            add_para(doc, "See per-branch detail tables below.", size=10, italic=True)
        doc.add_paragraph()

        # ── SECTION 4: PER-BRANCH DETAIL ──
        h1(doc, "4. Per-Branch Detail")
        for b in bdata:
            h2(doc, b["name"])
            rows = [
                ("Data points", f"{b['pts']} centerline points", HX_WHITE),
                (
                    "Z coordinate range",
                    f"{b['zTop']:.1f} → {b['zBot']:.1f} mm (Superior → Inferior)",
                    HX_WHITE,
                ),
                ("Centerline length", f"{b['length']:.1f} mm", HX_WHITE),
                ("Min diameter", f"{b['minD']:.2f} mm", HX_WHITE),
                (
                    "Max diameter",
                    f"{b['maxD']:.2f} mm",
                    HX_AMBER if b["maxD"] > 15 else HX_LGRN,
                ),
                ("Mean diameter", f"{b['avgD']:.2f} mm", HX_WHITE),
            ]
            # Ostium confidence row
            _oc = b.get("ost_conf")
            if _oc:
                _cg = _oc["grade"]
                _cs = _oc.get(
                    "effective_score", _oc["score"]
                )  # lower-bound score drives grade
                _cs_raw = _oc["score"]
                _cc = _oc["components"]
                _flags = ", ".join(_oc.get("flags", [])) or "—"
                _pen = _oc.get("flag_penalty", 0.0)
                _cbg = (
                    HX_LGRN
                    if _cg == "HIGH"
                    else (HX_AMBER if _cg in ("MEDIUM", "LOW") else HX_WHITE)
                )
                rows.append(
                    (
                        "Ostium confidence",
                        f"{_cg}  eff={_cs:.2f} (raw={_cs_raw:.2f} pen={_pen:.2f} ±{_oc.get('uncertainty',0):.2f})  "
                        f"V={_cc.get('V',0):.2f} A={_cc.get('A',0):.2f} "
                        f"L={_cc.get('L',0):.2f} T={_cc.get('T',0):.2f} "
                        f"G={_cc.get('G',0):.2f} C={_cc.get('C',0):.2f} "
                        f"Z={_cc.get('Z',0):.2f} D={_cc.get('D',0):.2f}  flags={_flags}",
                        _cbg,
                    )
                )
            rows.append(
                (
                    "Clinical note",
                    (
                        ai["findings"][:120] + "..."
                        if ai["findings"]
                        else "See AI summary."
                    ),
                    HX_WHITE,
                )
            )
            detail_table(doc, rows)
            doc.add_paragraph()

        # ── SECTION 5: STENT SIZING ──
        h1(doc, "5. Stent Sizing Recommendations")
        if ai["stent"]:
            add_para(doc, ai["stent"], size=10)
            doc.add_paragraph()

        st_w = [4.5, 2.0, 3.0, 6.0]
        st_hdr = ["Vessel", "Max Ø (mm)", "Suggested Stent", "Note"]
        st_tbl = doc.add_table(rows=1, cols=4)
        st_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (h, w) in enumerate(zip(st_hdr, st_w)):
            hdr_cell(st_tbl.rows[0].cells[i], h, w)

        stent_notes = {}  # no anatomy-specific notes
        for b in bdata:
            row = st_tbl.add_row()
            suggested = f"{int(b['maxD'])+1}–{int(b['maxD'])+3} mm diameter"
            note = stent_notes.get(b["name"], "Confirm with cross-section")
            note = note.format(avg=f"{b['avgD']:.1f}")
            row_vals = [b["name"], f"{b['maxD']:.2f}", suggested, note]
            for j, (txt, w) in enumerate(zip(row_vals, st_w)):
                c = row.cells[j]
                c.width = Cm(w)
                cell_fmt(
                    c,
                    txt,
                    bold=(j == 0),
                    size=10,
                    bg=HX_WHITE,
                    align=(
                        WD_ALIGN_PARAGRAPH.LEFT
                        if j in (0, 3)
                        else WD_ALIGN_PARAGRAPH.CENTER
                    ),
                )

        doc.add_paragraph()

        # ── SECTION 5b: FINDINGS TABLE ──
        # Combine pancaking + mild compression findings for report
        findings = getattr(self, "findings", [])
        # Also include mild compressions alongside pancaking/aneurysm/ectasia
        all_findings = findings  # already contains all types
        if all_findings:
            h1(doc, "5. Pathological Findings")
            findings = all_findings  # use full list in table below
            branch_labels_r = [self.getBranchDisplayName(i) for i in range(20)]
            type_colors = {
                "Aneurysm": ("FFE0E0", "8B0000"),
                "Ectasia": ("FFE0B2", "7D4E00"),
                "Pancaking": ("FFF9C4", "5D4037"),
                "Mild Compression": ("E8F5E9", "1B5E20"),
            }
            fw_cm = [2.5, 2.5, 1.8, 2.5, 6.2]
            fw_hdr = ["Type", "Vessel", "Z (mm)", "Distance (mm)", "Description"]
            f_tbl = doc.add_table(rows=1, cols=5)
            f_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i, (h, w) in enumerate(zip(fw_hdr, fw_cm)):
                hdr_cell(f_tbl.rows[0].cells[i], h, w)
            for f in findings:
                row = f_tbl.add_row()
                bg, tc = type_colors.get(f["type"], ("FFFFFF", "000000"))
                vessel = (
                    branch_labels_r[f["branchIdx"]]
                    if f["branchIdx"] < len(branch_labels_r)
                    else f"Branch {f['branchIdx']+1}"
                )
                vals = [f["type"], vessel, str(f["z"]), str(f["dist"]), f["desc"]]
                for j, (val, w) in enumerate(zip(vals, fw_cm)):
                    c = row.cells[j]
                    c.width = Cm(w)
                    cell_fmt(
                        c,
                        val,
                        bold=(j == 0),
                        color=RGBColor.from_string(tc),
                        size=10,
                        bg=bg,
                        align=(
                            WD_ALIGN_PARAGRAPH.LEFT
                            if j in (0, 1, 4)
                            else WD_ALIGN_PARAGRAPH.CENTER
                        ),
                    )
            doc.add_paragraph()

        # ── SECTION 5c: COLLATERALS TABLE ──
        collaterals = getattr(self, "collaterals", [])
        if collaterals:
            h1(doc, "6. Collateral Vessels")
            add_para(
                doc,
                "The following branches were identified as potential collateral vessels "
                "based on branch-off angle, diameter, and length criteria.",
                size=10,
                italic=True,
                color=C_GREY,
            )
            cw = [3.5, 2.0, 1.8, 1.8, 2.0, 1.8, 3.5]
            ch = [
                "Vessel",
                "Confidence",
                "Length mm",
                "Max Ø mm",
                "Angle °",
                "Origin Z",
                "Criteria",
            ]
            c_tbl = doc.add_table(rows=1, cols=7)
            c_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i, (h, w) in enumerate(zip(ch, cw)):
                hdr_cell(c_tbl.rows[0].cells[i], h, w)
            for c in collaterals:
                crow = c_tbl.add_row()
                conf_bg = "E0F7FA" if c["confidence"] == "High" else "F5F5F5"
                cvals = [
                    c["label"],
                    c["confidence"],
                    f"{c['length_mm']:.0f}",
                    f"{c['maxD']:.1f}",
                    f"{c['angle_deg']:.0f}",
                    f"{c['originZ']:.0f}",
                    c["reasons"],
                ]
                for j, (val, w) in enumerate(zip(cvals, cw)):
                    cc = crow.cells[j]
                    cc.width = Cm(w)
                    cell_fmt(
                        cc,
                        val,
                        bold=(j == 0),
                        size=10,
                        bg=conf_bg if j < 2 else HX_WHITE,
                        color=RGBColor(0x00, 0x60, 0x64) if j == 1 else C_BLACK,
                        align=(
                            WD_ALIGN_PARAGRAPH.LEFT
                            if j in (0, 6)
                            else WD_ALIGN_PARAGRAPH.CENTER
                        ),
                    )
            doc.add_paragraph()

        # ── SECTION 6: LIMITATIONS ──
        h1(doc, "7. Limitations & Next Steps")
        if ai["limitations"]:
            add_para(doc, ai["limitations"], size=10)
        lims = [
            "Diameter computed via closest surface point — may underestimate in compressed vessels.",
            "Min diameter values near vessel endpoints may be artifacts — verify before clinical use.",
            "All stent sizing decisions must be confirmed by the treating physician with cross-sectional imaging.",
            "This report was generated with AI assistance and has not been reviewed by a radiologist.",
        ]
        for lim in lims:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"•  {lim}")
            run.font.size = Pt(10)
            run.font.color.rgb = C_BLACK

        doc.add_paragraph()

        # ── STENT PLANNING SECTION ───────────────────────────────────────
        plan = getattr(self, "stentPlan", None)
        if plan:
            h1(doc, "8. Stent Planning")
            stent_rows = [
                ("Stent type", plan["type"]),
                (
                    "Proximal diameter",
                    str(plan["proxDiam"])
                    + " mm  (vessel: "
                    + str(plan["vd_prox"])
                    + " mm)",
                ),
                (
                    "Distal diameter",
                    str(plan["distDiam"])
                    + " mm  (vessel: "
                    + str(plan["vd_dist"])
                    + " mm)",
                ),
                ("Planned length", str(plan["length"]) + " mm"),
                ("Proximal landing pt", "Point " + str(plan["proxPt"])),
                ("Distal landing pt", "Point " + str(plan["distPt"])),
                (
                    "Sizing status",
                    "OK" if not plan["warnings"] else "; ".join(plan["warnings"]),
                ),
            ]
            st_tbl = doc.add_table(rows=len(stent_rows), cols=2)
            st_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for ri, (lbl, val) in enumerate(stent_rows):
                row = st_tbl.rows[ri]
                is_warn = ri == len(stent_rows) - 1 and plan["warnings"]
                bg = "FFE0CC" if is_warn else ("F0F4FF" if ri % 2 == 0 else "FFFFFF")
                cell_fmt(
                    row.cells[0],
                    lbl,
                    bold=True,
                    size=10,
                    bg=bg,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )
                cell_fmt(
                    row.cells[1],
                    val,
                    bold=False,
                    size=10,
                    bg=bg,
                    color=RGBColor(0xC0, 0x39, 0x2B) if is_warn else C_BLACK,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )
            doc.add_paragraph()

        # ── SECTION 7: CONCLUSION ──
        h1(doc, "7. Conclusion")
        if ai["conclusion"]:
            add_para(doc, ai["conclusion"], size=10)
        doc.add_paragraph()

        # ── END ──
        ep = doc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        er = ep.add_run("— END OF REPORT —")
        er.italic = True
        er.font.size = Pt(9)
        er.font.color.rgb = C_SILVER

        foot = doc.add_paragraph()
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = foot.add_run(
            "Generated by VesselAnalyzer (3D Slicer / VMTK) with AI assistance (Claude Sonnet). "
            "For clinical review only."
        )
        fr.italic = True
        fr.font.size = Pt(8)
        fr.font.color.rgb = C_SILVER

        doc.save(path)

        # Fix internal docx metadata date (python-docx template has 2013 date baked in)
        import zipfile, shutil, os
        from datetime import datetime, timezone

        now_iso = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        tmp_path = path + ".tmp.docx"
        with zipfile.ZipFile(path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "docProps/core.xml":
                        import re

                        data = data.decode("utf-8")
                        data = re.sub(
                            r"<dcterms:created[^>]*>.*?</dcterms:created>",
                            f'<dcterms:created xsi:type="dcterms:W3CDTF">{now_iso}</dcterms:created>',
                            data,
                        )
                        data = re.sub(
                            r"<dcterms:modified[^>]*>.*?</dcterms:modified>",
                            f'<dcterms:modified xsi:type="dcterms:W3CDTF">{now_iso}</dcterms:modified>',
                            data,
                        )
                        data = data.encode("utf-8")
                    zout.writestr(item, data)
        os.replace(tmp_path, path)
        print(f"[VesselAnalyzer] AI report saved: {path}")


    def exportToCSV(self, path):
        with open(path, "w") as f:
            f.write("PointIndex,X,Y,Z,DistanceFromStart_mm,Diameter_mm,Radius_mm\n")
            for i, pt in enumerate(self.points):
                d = self.diameters[i] if self.diameters else 0.0
                f.write(f"{i},{pt[0]:.4f},{pt[1]:.4f},{pt[2]:.4f},{d:.4f}\n")


# ── Slicer module-scanner guard ───────────────────────────────────────────────
# Slicer auto-scans all .py files in the module folder and expects a class
# matching the filename.  This stub satisfies that requirement without
# registering as a real loadable module (no ScriptedLoadableModule base).
class vessel_report_mixin:  # noqa: E302
    """Slicer module-scanner stub — not a real loadable module."""
    def __init__(self, parent=None):
        if parent:
            parent.title = "vessel_report_mixin"
            parent.hidden = True  # hide from Slicer module list
